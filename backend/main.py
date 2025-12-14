"""
MingJing - 中文敏感信息识别系统 API

FastAPI 后端服务
"""

import os
import re
import sys
import tempfile
import time
from typing import List, Optional, Dict, Any
from contextlib import asynccontextmanager
from pathlib import Path
import io
import zipfile

from fastapi import FastAPI, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, Response
from pydantic import BaseModel, Field

# 添加模块路径
sys.path.insert(0, os.path.dirname(__file__))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "presidio-analyzer"))

from config import get_settings
from config.rules_manager import get_rules_manager, Rule, PatternConfig
from core.providers import AnalyzerEngineProvider
from core.processors import FileProcessor
from core.reporters import ReportGenerator, FileResult, EntityInfo
from core.metrics import get_metrics_collector
from db import get_history_db


# ==================== 数据模型 ====================

class AnalyzeRequest(BaseModel):
    text: str = Field(..., description="待分析的文本", min_length=1)
    entities: Optional[List[str]] = Field(None, description="要识别的实体类型")
    score_threshold: float = Field(0.0, ge=0.0, le=1.0)
    use_llm_verify: bool = Field(False)


class EntityResult(BaseModel):
    entity_type: str
    text: str
    start: int
    end: int
    score: float
    verified: bool = False
    llm_reason: Optional[str] = None


class AnalyzeResponse(BaseModel):
    text: str
    results: List[EntityResult]
    count: int


class EntityTypeInfo(BaseModel):
    name: str
    description: str
    category: str


class HealthResponse(BaseModel):
    status: str
    nlp_loaded: bool
    recognizers_count: int
    version: str


class RuleRequest(BaseModel):
    name: str
    entity_type: str
    patterns: List[Dict[str, Any]]
    context: List[str] = []
    enabled: bool = True
    description: str = ""
    category: str = "其他"


class ConfigUpdateRequest(BaseModel):
    key: str
    value: Any


class RecognizerUpdateRequest(BaseModel):
    enabled: Optional[bool] = None
    base_score: Optional[float] = Field(None, ge=0.0, le=1.0)
    description: Optional[str] = None


class ContextWordsUpdateRequest(BaseModel):
    words: List[str]


# ==================== 全局变量 ====================

engine_provider: Optional[AnalyzerEngineProvider] = None
file_processor: Optional[FileProcessor] = None
supported_entities: List[EntityTypeInfo] = []


# ==================== 实体类型描述 ====================

ENTITY_DESCRIPTIONS = {
    "CN_ID_CARD": ("身份证号", "个人身份"),
    "CN_PHONE": ("手机号/电话", "联系方式"),
    "CN_BANK_CARD": ("银行卡号", "金融信息"),
    "CN_EMAIL": ("电子邮箱", "联系方式"),
    "CN_IP_ADDRESS": ("IP地址", "网络信息"),
    "CN_POSTAL_CODE": ("邮政编码", "地址信息"),
    "CN_VEHICLE_PLATE": ("车牌号", "车辆信息"),
    "CN_PASSPORT": ("护照号", "个人身份"),
    "CN_DRIVER_LICENSE": ("驾驶证号", "个人身份"),
    "CN_MILITARY_ID": ("军人证号", "个人身份"),
    "CN_SOCIAL_CREDIT_CODE": ("统一社会信用代码", "企业信息"),
    "CN_MEDICAL_LICENSE": ("医师执业证号", "医疗信息"),
    "CN_MAC_ADDRESS": ("MAC地址", "网络信息"),
    "CN_JDBC": ("JDBC连接字符串", "网络信息"),
    "CN_JWT": ("JWT Token", "认证信息"),
    "CN_CLOUD_KEY": ("云服务密钥", "认证信息"),
    "CN_WECHAT": ("微信ID", "联系方式"),
    "CN_SENSITIVE_FIELD": ("敏感字段", "认证信息"),
    "PERSON": ("人名", "NLP识别"),
    "LOCATION": ("地名/地址", "NLP识别"),
    "ORGANIZATION": ("组织机构", "NLP识别"),
}


# ==================== 初始化 ====================

def init_engine():
    """初始化分析引擎"""
    global engine_provider, file_processor, supported_entities
    
    settings = get_settings(reload=True)
    
    # 强制重新加载规则管理器
    get_rules_manager(reload=True)
    
    # 创建分析引擎
    engine_provider = AnalyzerEngineProvider(settings)
    engine_provider.create()
    
    # 创建文件处理器
    file_processor = FileProcessor(
        max_file_size=settings.file_processing.max_file_size * 1024 * 1024,
        temp_dir=settings.file_processing.temp_dir,
    )
    
    # 获取支持的实体类型
    entities = engine_provider.get_supported_entities()
    supported_entities.clear()
    for entity in entities:
        desc, category = ENTITY_DESCRIPTIONS.get(entity, (entity, "其他"))
        supported_entities.append(EntityTypeInfo(name=entity, description=desc, category=category))
    
    print(f"✅ 分析引擎初始化完成，共 {len(supported_entities)} 种实体类型")


@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    init_engine()
    yield
    # 清理
    if file_processor:
        file_processor.cleanup()


# ==================== FastAPI 应用 ====================

app = FastAPI(
    title="MingJing - 中文敏感信息识别系统",
    description="基于 Presidio 的中文敏感信息识别 API",
    version="1.0.0",
    lifespan=lifespan,
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==================== 系统接口 ====================

@app.get("/health", response_model=HealthResponse, tags=["系统"])
async def health_check():
    """健康检查"""
    settings = get_settings()
    return HealthResponse(
        status="healthy",
        nlp_loaded=engine_provider.nlp_engine is not None if engine_provider else False,
        recognizers_count=len(supported_entities),
        version=settings.version,
    )


@app.get("/entities", response_model=List[EntityTypeInfo], tags=["系统"])
async def get_supported_entities():
    """获取支持的实体类型"""
    return supported_entities


@app.post("/reload", tags=["系统"])
async def reload_engine():
    """重新加载分析引擎"""
    try:
        init_engine()
        return {"message": "重新加载成功", "entities_count": len(supported_entities)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== 结果过滤 ====================

# 常见误报词汇（不应被识别为人名）
FALSE_POSITIVE_NAMES = {
    # 英文产品/技术名
    'github', 'slack', 'docker', 'redis', 'mysql', 'nginx', 'apache',
    'linux', 'windows', 'macos', 'python', 'java', 'golang', 'rust',
    'kubernetes', 'jenkins', 'gitlab', 'bitbucket', 'aws', 'azure', 'gcp',
    'ghp', 'gho', 'ghs', 'ghu', 'xoxb', 'xoxp',  # token前缀
    # 常见单字误报
    '赵', '钱', '孙', '李', '周', '吴', '郑', '王', '冯', '陈',
}

# 日期时间误报模式
FALSE_POSITIVE_DATETIME_PATTERNS = [
    r'^[a-f0-9]{20,}$',  # 纯hex字符串
    r'^[A-Za-z0-9]{30,}$',  # 长token
]

def filter_false_positives(results: list, text: str) -> list:
    """过滤明显的误报结果"""
    filtered = []
    for r in results:
        entity_text = text[r.start:r.end] if hasattr(r, 'start') else r.text
        entity_type = r.entity_type
        
        # 过滤人名误报
        if entity_type == 'PERSON':
            # 过滤单字
            if len(entity_text) <= 1:
                continue
            # 过滤常见误报词
            if entity_text.lower() in FALSE_POSITIVE_NAMES:
                continue
            # 过滤纯英文小写（可能是变量名）
            if entity_text.isascii() and entity_text.islower() and len(entity_text) < 5:
                continue
        
        # 完全移除日期时间实体
        if entity_type == 'DATE_TIME':
            continue
        
        filtered.append(r)
    
    return filtered


# ==================== 脱敏工具 ====================

def _mask_value(entity_type: str, value: str) -> str:
    """智能脱敏策略：根据实体类型保留部分信息。"""
    if not value:
        return value
    
    length = len(value)
    
    if entity_type == 'CN_PHONE':
        # 手机号：138****5678 (保留前3后4)
        if length == 11:
            return value[:3] + '****' + value[7:]
        return value[:min(3, length)] + '****'
    
    elif entity_type == 'CN_ID_CARD':
        # 身份证：3301**********1234 (保留前4后4)
        if length == 18:
            return value[:4] + '*' * 10 + value[14:]
        elif length == 15:
            return value[:4] + '*' * 7 + value[11:]
        return value[:4] + '*' * max(0, length - 8) + value[max(4, length - 4):]
    
    elif entity_type == 'CN_BANK_CARD':
        # 银行卡：6222 **** **** 1234 (保留前4后4)
        if length >= 16:
            return value[:4] + ' **** **** ' + value[-4:]
        return value[:4] + '****' + value[max(4, length - 4):]
    
    elif entity_type == 'CN_EMAIL':
        # 邮箱：abc***@example.com (保留前3和域名)
        at_index = value.find('@')
        if at_index > 0:
            local_part = value[:at_index]
            domain = value[at_index:]
            if len(local_part) <= 3:
                return local_part[0] + '***' + domain
            return local_part[:3] + '***' + domain
        return value[:3] + '***'
    
    elif entity_type == 'PERSON':
        # 姓名：张* 或 欧阳** (保留姓氏)
        if length == 2:
            return value[0] + '*'
        elif length == 3:
            return value[0] + '**'
        elif length >= 4:
            # 复姓情况
            return value[:2] + '*' * (length - 2)
        return value[0] + '*'
    
    elif entity_type in ('CN_PASSPORT', 'CN_DRIVER_LICENSE', 'CN_MILITARY_ID'):
        # 证件号：保留前2后2
        if length > 4:
            return value[:2] + '*' * (length - 4) + value[-2:]
        return '***'
    
    elif entity_type == 'CN_VEHICLE_PLATE':
        # 车牌：京A****8 (保留省份+首字母和最后1位)
        if length >= 7:
            return value[:2] + '****' + value[-1:]
        return value[:2] + '****'
    
    elif entity_type == 'LOCATION':
        # 地址：保留前6个字
        if length > 10:
            return value[:6] + '***'
        return value[:min(4, length)] + '***'
    
    else:
        # 其他类型：全部替换为黑圆点
        return '●' * length


def _apply_masks_to_text(text: str, results: list) -> str:
    """根据识别结果在原文本上应用掩码。"""
    if not results:
        return text

    # 按start排序，避免重叠冲突
    sorted_results = sorted(results, key=lambda r: r.start)
    parts = []
    last_end = 0

    for r in sorted_results:
        start, end = r.start, r.end
        # 跳过与前一个结果重叠的部分
        if start < last_end:
            continue
        parts.append(text[last_end:start])
        original = text[start:end]
        parts.append(_mask_value(r.entity_type, original))
        last_end = end

    parts.append(text[last_end:])
    return "".join(parts)


def anonymize_text_segment(text: str, entities: Optional[List[str]], score_threshold: float) -> str:
    """对一段纯文本做识别+脱敏。"""
    if not text or not engine_provider or not engine_provider.analyzer:
        return text

    results = engine_provider.analyzer.analyze(
        text=text,
        language="zh",
        entities=entities,
    )

    # 按阈值和误报过滤
    results = [r for r in results if r.score >= score_threshold]
    results = filter_false_positives(results, text)

    return _apply_masks_to_text(text, results)


def anonymize_text_bytes(data: bytes, filename: str, entities: Optional[List[str]], score_threshold: float) -> bytes:
    """对纯文本类文件（txt/csv/json等）做脱敏，保持整体结构。"""
    try:
        text = data.decode("utf-8", errors="ignore")
    except Exception:
        # 解码失败直接返回原文
        return data

    redacted = anonymize_text_segment(text, entities, score_threshold)
    return redacted.encode("utf-8")


def anonymize_docx_bytes(data: bytes, entities: Optional[List[str]], score_threshold: float) -> bytes:
    """对 DOCX 文档做脱敏，尽量保持结构。"""
    try:
        from docx import Document
    except ImportError:
        # 环境不支持时直接返回原文
        return data

    file_obj = io.BytesIO(data)
    doc = Document(file_obj)

    # 段落级脱敏（按整段文本处理，再回写）
    for para in doc.paragraphs:
        if para.text:
            redacted = anonymize_text_segment(para.text, entities, score_threshold)
            if redacted != para.text:
                para.text = redacted

    # 表格单元格脱敏
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    redacted = anonymize_text_segment(cell.text, entities, score_threshold)
                    if redacted != cell.text:
                        cell.text = redacted

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def anonymize_xlsx_bytes(data: bytes, entities: Optional[List[str]], score_threshold: float) -> bytes:
    """对 XLSX/XLS 文档做脱敏，只修改文本单元格。"""
    try:
        import openpyxl
    except ImportError:
        return data

    file_obj = io.BytesIO(data)
    try:
        wb = openpyxl.load_workbook(file_obj, data_only=False)
    except Exception:
        return data

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                val = cell.value
                if isinstance(val, str) and val:
                    redacted = anonymize_text_segment(val, entities, score_threshold)
                    if redacted != val:
                        cell.value = redacted

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def anonymize_zip_bytes(data: bytes, entities: Optional[List[str]], score_threshold: float) -> bytes:
    """对 ZIP 中的支持类型文件做递归脱敏并重新打包。"""
    in_io = io.BytesIO(data)
    out_io = io.BytesIO()

    try:
        with zipfile.ZipFile(in_io, "r") as zin, zipfile.ZipFile(out_io, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for info in zin.infolist():
                file_data = zin.read(info.filename)

                # 目录条目直接写入
                if info.is_dir():
                    zout.writestr(info, b"")
                    continue

                ext = Path(info.filename).suffix.lower()

                if ext in {".txt", ".log", ".md", ".csv", ".tsv", ".json", ".jsonl"}:
                    redacted = anonymize_text_bytes(file_data, info.filename, entities, score_threshold)
                elif ext == ".docx":
                    redacted = anonymize_docx_bytes(file_data, entities, score_threshold)
                elif ext in {".xlsx", ".xls"}:
                    redacted = anonymize_xlsx_bytes(file_data, entities, score_threshold)
                else:
                    # 其他类型原样保留
                    redacted = file_data

                zout.writestr(info, redacted)
    except Exception:
        # 任意异常回退为原数据
        return data

    out_io.seek(0)
    return out_io.getvalue()


# ==================== 分析接口 ====================

@app.post("/analyze", response_model=AnalyzeResponse, tags=["分析"])
async def analyze_text(request: AnalyzeRequest):
    """分析文本中的敏感信息"""
    if not engine_provider or not engine_provider.analyzer:
        raise HTTPException(status_code=503, detail="分析引擎未初始化")
    
    collector = get_metrics_collector()
    
    with collector.track_request():
        # 执行分析
        results = engine_provider.analyzer.analyze(
            text=request.text,
            language="zh",
            entities=request.entities,
        )
        
        # 过滤低分结果
        results = [r for r in results if r.score >= request.score_threshold]
        
        # 过滤误报
        results = filter_false_positives(results, request.text)
        
        # LLM验证
        entity_results = []
        if request.use_llm_verify and engine_provider.verifier:
            verified = engine_provider.verifier.verify_results(request.text, results)
            for original, verification in verified:
                entity_text = request.text[original.start:original.end]
                if verification:
                    if verification.is_sensitive:
                        entity_results.append(EntityResult(
                            entity_type=original.entity_type,
                            text=entity_text,
                            start=original.start,
                            end=original.end,
                            score=verification.final_score,
                            verified=True,
                            llm_reason=verification.reason,
                        ))
                        collector.record_entity(original.entity_type)
                else:
                    entity_results.append(EntityResult(
                        entity_type=original.entity_type,
                        text=entity_text,
                        start=original.start,
                        end=original.end,
                        score=original.score,
                    ))
                    collector.record_entity(original.entity_type)
        else:
            for r in results:
                entity_results.append(EntityResult(
                    entity_type=r.entity_type,
                    text=request.text[r.start:r.end],
                    start=r.start,
                    end=r.end,
                    score=r.score,
                ))
                collector.record_entity(r.entity_type)
    
    return AnalyzeResponse(
        text=request.text,
        results=entity_results,
        count=len(entity_results),
    )


# 文件大小限制：50MB
MAX_FILE_SIZE = 50 * 1024 * 1024

@app.post("/analyze/file", tags=["分析"])
async def analyze_file(
    file: UploadFile = File(...),
    entities: Optional[str] = Form(None),
    score_threshold: float = Form(0.0),
):
    """分析上传的文件"""
    if not engine_provider or not file_processor:
        raise HTTPException(status_code=503, detail="服务未初始化")
    
    start_time = time.time()
    
    # 读取文件
    content = await file.read()
    
    # 检查文件大小
    if len(content) > MAX_FILE_SIZE:
        size_mb = len(content) / 1024 / 1024
        raise HTTPException(
            status_code=413,
            detail=f"文件过大 ({size_mb:.2f}MB)，最大支持 50MB"
        )
    
    # 处理文件
    processed = file_processor.process_bytes(content, file.filename)
    
    if processed.error:
        raise HTTPException(status_code=400, detail=processed.error)
    
    # 分析内容
    entity_list = entities.split(",") if entities else None
    all_results = []
    
    for extracted in processed.contents:
        results = engine_provider.analyzer.analyze(
            text=extracted.text,
            language="zh",
            entities=entity_list,
        )
        
        # 过滤低分和误报
        results = [r for r in results if r.score >= score_threshold]
        results = filter_false_positives(results, extracted.text)
        
        for r in results:
            all_results.append({
                "entity_type": r.entity_type,
                "text": extracted.text[r.start:r.end],
                "start": r.start,
                "end": r.end,
                "score": r.score,
                "source": extracted.source,
            })
    
    process_time = time.time() - start_time
    
    return {
        "filename": file.filename,
        "file_type": processed.file_type,
        "file_size": processed.size,
        "is_archive": processed.is_archive,
        "content_blocks": processed.content_count,
        "results": all_results,
        "count": len(all_results),
        "process_time": round(process_time, 2),
    }


@app.post("/anonymize/file", tags=["脱敏"])
async def anonymize_file(
    file: UploadFile = File(...),
    entities: Optional[str] = Form(None),
    score_threshold: float = Form(0.0),
):
    """对上传的文件进行自动脱敏并返回脱敏后的文件。

    当前支持: TXT/CSV/JSON/DOCX/XLSX/ZIP。
    """
    if not engine_provider:
        raise HTTPException(status_code=503, detail="服务未初始化")

    data = await file.read()
    
    # 检查文件大小
    if len(data) > MAX_FILE_SIZE:
        size_mb = len(data) / 1024 / 1024
        raise HTTPException(
            status_code=413,
            detail=f"文件过大 ({size_mb:.2f}MB)，最大支持 50MB"
        )
    ext = Path(file.filename).suffix.lower()
    entity_list = entities.split(",") if entities else None

    if ext in {".txt", ".log", ".md", ".csv", ".tsv", ".json", ".jsonl"}:
        redacted = anonymize_text_bytes(data, file.filename, entity_list, score_threshold)
    elif ext == ".docx":
        redacted = anonymize_docx_bytes(data, entity_list, score_threshold)
    elif ext in {".xlsx", ".xls"}:
        redacted = anonymize_xlsx_bytes(data, entity_list, score_threshold)
    elif ext == ".zip":
        redacted = anonymize_zip_bytes(data, entity_list, score_threshold)
    else:
        raise HTTPException(status_code=400, detail=f"暂不支持该文件类型脱敏: {ext}")

    # 根据扩展名设置大致的 MIME 类型
    if ext in {".txt", ".log", ".md"}:
        media_type = "text/plain; charset=utf-8"
    elif ext in {".csv", ".tsv"}:
        media_type = "text/csv; charset=utf-8"
    elif ext in {".json", ".jsonl"}:
        media_type = "application/json; charset=utf-8"
    elif ext == ".docx":
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif ext in {".xlsx", ".xls"}:
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif ext == ".zip":
        media_type = "application/zip"
    else:
        media_type = "application/octet-stream"

    return Response(
        content=redacted,
        media_type=media_type,
        headers={
            "Content-Disposition": f'attachment; filename="{file.filename}"',
        },
    )


# ==================== 规则管理接口 ====================

@app.get("/rules", tags=["规则管理"])
async def get_rules(source: Optional[str] = None):
    """
    获取所有规则
    
    :param source: 可选过滤，builtin=内置规则，custom=自定义规则，不传=全部
    """
    manager = get_rules_manager()
    
    if source == "builtin":
        rules = manager.get_builtin_rules()
    elif source == "custom":
        rules = manager.get_custom_rules()
    else:
        rules = manager.get_all_rules()
    
    return [rule.to_dict() for rule in rules]


@app.post("/rules", tags=["规则管理"])
async def add_rule(request: RuleRequest):
    """添加自定义规则（内置规则不能通过此接口添加）"""
    manager = get_rules_manager()
    
    patterns = [
        PatternConfig(
            regex=p.get("regex", ""),
            name=p.get("name", ""),
            score=p.get("score", 0.5),
        )
        for p in request.patterns
    ]
    
    rule = Rule(
        name=request.name,
        entity_type=request.entity_type,
        description=request.description,
        category=request.category,
        source="custom",  # 新增规则只能是自定义
        patterns=patterns,
        context=request.context,
        enabled=request.enabled,
    )
    
    success, message = manager.add_rule(rule)
    if not success:
        raise HTTPException(status_code=400, detail=message)
    
    manager.save()
    init_engine()  # 重新加载引擎
    return {"message": message}


@app.put("/rules/{name}", tags=["规则管理"])
async def update_rule(name: str, request: RuleRequest):
    """
    更新规则（内置和自定义都可以修改）
    
    内置规则可以修改：正则、分数、上下文、启用状态
    内置规则不能：删除、改名
    """
    manager = get_rules_manager()
    
    patterns = [
        PatternConfig(
            regex=p.get("regex", ""),
            name=p.get("name", ""),
            score=p.get("score", 0.5),
        )
        for p in request.patterns
    ]
    
    # 获取原规则以保持 source
    old_rule = manager.get_rule(name)
    if not old_rule:
        raise HTTPException(status_code=404, detail=f"规则 '{name}' 不存在")
    
    rule = Rule(
        name=request.name,
        entity_type=request.entity_type,
        description=request.description,
        category=request.category,
        source=old_rule.source,  # 保持原来的来源
        patterns=patterns,
        context=request.context,
        enabled=request.enabled,
    )
    
    success, message = manager.update_rule(name, rule)
    if not success:
        raise HTTPException(status_code=400, detail=message)
    
    manager.save()
    init_engine()  # 重新加载引擎
    return {"message": message, "rule": rule.to_dict()}


@app.delete("/rules/{name}", tags=["规则管理"])
async def delete_rule(name: str):
    """删除规则（仅自定义规则可删除，内置规则不能删除）"""
    manager = get_rules_manager()
    success, message = manager.delete_rule(name)
    if not success:
        raise HTTPException(status_code=400, detail=message)
    
    manager.save()
    init_engine()  # 重新加载引擎
    return {"message": message}


@app.post("/rules/{name}/test", tags=["规则管理"])
async def test_rule(name: str, text: str = Form(...)):
    """测试规则"""
    manager = get_rules_manager()
    results = manager.test_rule(name, text)
    return {"results": results}


@app.post("/rules/validate-regex", tags=["规则管理"])
async def validate_regex(regex: str = Form(...)):
    """验证正则表达式"""
    manager = get_rules_manager()
    valid, message = manager.validate_regex(regex)
    return {"valid": valid, "message": message}


# ==================== 配置接口 ====================

@app.get("/config", tags=["配置"])
async def get_config():
    """获取当前配置"""
    settings = get_settings()
    return settings.to_dict()


@app.put("/config", tags=["配置"])
async def update_config(request: ConfigUpdateRequest):
    """更新配置"""
    settings = get_settings()
    try:
        settings.set(request.key, request.value)
        return {"message": f"配置 {request.key} 已更新"}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/recognizers/{name}", tags=["配置"])
async def update_recognizer(name: str, request: RecognizerUpdateRequest):
    """更新内置识别器配置"""
    settings = get_settings()

    if (
        request.base_score is None
        and request.enabled is None
        and request.description is None
    ):
        raise HTTPException(status_code=400, detail="未提供任何可更新的字段")

    try:
        settings.update_recognizer(
            name,
            enabled=request.enabled,
            base_score=request.base_score,
            description=request.description,
        )
        init_engine()
        return {
            "message": f"识别器 {name} 已更新",
            "settings": settings.recognizers.settings.get(name). __dict__ if name in settings.recognizers.settings else None,
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.put("/context/{entity_type}", tags=["配置"])
async def update_context(entity_type: str, request: ContextWordsUpdateRequest):
    """更新实体上下文词"""
    settings = get_settings()
    cleaned = [word.strip() for word in request.words if word.strip()]

    try:
        settings.update_context_words(entity_type, cleaned)
        init_engine()
        return {"message": f"实体 {entity_type} 的上下文词已更新", "words": cleaned}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


# ==================== 报告接口 ====================

@app.post("/report/generate", tags=["报告"])
async def generate_report(
    title: str = Form("敏感信息识别报告"),
    format: str = Form("json"),
    background_tasks: BackgroundTasks = None,
):
    """生成报告（需要先进行分析）"""
    # 这里需要实现报告生成逻辑
    # 暂时返回示例
    return {"message": "报告生成功能开发中"}


# ==================== 性能监控接口 ====================

@app.get("/metrics", tags=["监控"])
async def get_metrics():
    """获取性能指标"""
    collector = get_metrics_collector()
    return collector.get_metrics()


@app.post("/metrics/reset", tags=["监控"])
async def reset_metrics():
    """重置性能指标"""
    collector = get_metrics_collector()
    collector.reset()
    return {"message": "指标已重置"}


# ==================== LLM 配置接口 ====================

@app.get("/llm/status", tags=["LLM"])
async def get_llm_status():
    """获取 LLM 验证器状态"""
    from core.llm_integration import get_llm_manager
    manager = get_llm_manager()
    return manager.get_status()


@app.post("/llm/configure/local", tags=["LLM"])
async def configure_local_llm(
    model_path: str = Form(..., description="模型路径或HuggingFace模型名"),
    device: str = Form("auto", description="设备: auto/cpu/cuda"),
    max_new_tokens: int = Form(200, description="最大生成token数"),
):
    """配置本地 LLM 模型"""
    from core.llm_integration import get_llm_manager
    manager = get_llm_manager()
    
    success = manager.configure_local_model(
        model_path=model_path,
        device=device,
        max_new_tokens=max_new_tokens,
    )
    
    if success:
        # 更新全局验证器
        global engine_provider
        if engine_provider:
            engine_provider._verifier = manager.get_verifier()
        
        return {"message": "本地模型配置成功", "status": manager.get_status()}
    else:
        raise HTTPException(status_code=400, detail="配置失败")


@app.post("/llm/configure/api", tags=["LLM"])
async def configure_api_llm(
    api_key: str = Form(..., description="API 密钥"),
    api_base: str = Form("https://api.openai.com/v1", description="API 基础 URL"),
    model: str = Form("gpt-3.5-turbo", description="模型名称"),
    timeout: int = Form(30, description="超时时间(秒)"),
):
    """配置 API 模式 LLM"""
    from core.llm_integration import get_llm_manager
    manager = get_llm_manager()
    
    success = manager.configure_api(
        api_key=api_key,
        api_base=api_base,
        model=model,
        timeout=timeout,
    )
    
    if success:
        global engine_provider
        if engine_provider:
            engine_provider._verifier = manager.get_verifier()
        
        return {"message": "API 配置成功", "status": manager.get_status()}
    else:
        raise HTTPException(status_code=400, detail="配置失败")


@app.post("/llm/configure/mock", tags=["LLM"])
async def configure_mock_llm(
    default_is_sensitive: bool = Form(True),
    default_confidence: float = Form(0.8),
):
    """配置 Mock 模式（测试用）"""
    from core.llm_integration import get_llm_manager
    manager = get_llm_manager()
    
    manager.configure_mock(
        default_is_sensitive=default_is_sensitive,
        default_confidence=default_confidence,
    )
    
    global engine_provider
    if engine_provider:
        engine_provider._verifier = manager.get_verifier()
    
    return {"message": "Mock 模式配置成功", "status": manager.get_status()}


@app.post("/llm/reload", tags=["LLM"])
async def reload_llm():
    """重新加载 LLM 验证器"""
    from core.llm_integration import get_llm_manager
    manager = get_llm_manager()
    manager.reload()
    
    global engine_provider
    if engine_provider:
        engine_provider._verifier = manager.get_verifier()
    
    return {"message": "LLM 验证器已重新加载", "status": manager.get_status()}


# ==================== 历史记录 API ====================

@app.get("/history", tags=["历史记录"])
async def get_history(limit: int = 100):
    """获取历史记录"""
    db = get_history_db()
    records = db.get_all(limit=limit)
    return [r.to_dict() for r in records]


@app.post("/history", tags=["历史记录"])
async def add_history(
    text: str = Form(...),
    results: str = Form(...),  # JSON string
    record_type: str = Form("text"),
    filename: Optional[str] = Form(None),
    operation_type: str = Form("analyze"),  # 'analyze' or 'anonymize'
):
    """添加历史记录"""
    import json
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info(f"接收历史记录: record_type={record_type}, operation_type={operation_type}, filename={filename}")
    
    try:
        results_list = json.loads(results)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="results 必须是有效的 JSON")
    
    db = get_history_db()
    record = db.add_or_update(text, results_list, record_type, filename, operation_type)
    logger.info(f"保存成功: id={record.id}, operation_type={record.operation_type}")
    return record.to_dict()


@app.delete("/history/{record_id}", tags=["历史记录"])
async def delete_history(record_id: int):
    """删除指定历史记录"""
    db = get_history_db()
    if db.delete(record_id):
        return {"message": "删除成功"}
    raise HTTPException(status_code=404, detail="记录不存在")


@app.delete("/history", tags=["历史记录"])
async def clear_history():
    """清空所有历史记录"""
    db = get_history_db()
    count = db.clear_all()
    return {"message": f"已清空 {count} 条记录"}


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
