"""
文件内容提取器

将各种格式的文件转换为纯文本
"""

import csv
import json
import logging
from abc import ABC, abstractmethod
from io import StringIO, BytesIO
from pathlib import Path
from typing import Generator, Optional, Union
from dataclasses import dataclass

logger = logging.getLogger("mingjing.extractors")


@dataclass
class ExtractedContent:
    """提取的内容"""
    text: str
    source: str  # 来源（文件名、sheet名等）
    page: Optional[int] = None  # 页码（PDF等）
    row: Optional[int] = None  # 行号（CSV等）


class BaseExtractor(ABC):
    """提取器基类"""
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS: list = []
    
    @classmethod
    def supports(cls, file_path: Union[str, Path]) -> bool:
        """检查是否支持该文件类型"""
        ext = Path(file_path).suffix.lower()
        return ext in cls.SUPPORTED_EXTENSIONS
    
    @abstractmethod
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        """
        提取文件内容
        
        :param file_path: 文件路径
        :yield: ExtractedContent对象
        """
        pass
    
    @abstractmethod
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        """
        从字节数据提取内容
        
        :param data: 文件字节数据
        :param filename: 文件名
        :yield: ExtractedContent对象
        """
        pass


class TextExtractor(BaseExtractor):
    """纯文本提取器"""
    
    SUPPORTED_EXTENSIONS = [".txt", ".log", ".md", ".rst", ".ini", ".conf", ".cfg"]
    
    def __init__(self, encoding: str = "utf-8", chunk_size: int = 1024 * 1024):
        """
        初始化
        
        :param encoding: 文件编码
        :param chunk_size: 分块大小（字节）
        """
        self.encoding = encoding
        self.chunk_size = chunk_size
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            with open(path, 'r', encoding=self.encoding, errors='ignore') as f:
                chunk_num = 0
                while True:
                    chunk = f.read(self.chunk_size)
                    if not chunk:
                        break
                    yield ExtractedContent(
                        text=chunk,
                        source=path.name,
                        page=chunk_num,
                    )
                    chunk_num += 1
        except Exception as e:
            logger.error(f"读取文件失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            text = data.decode(self.encoding, errors='ignore')
            yield ExtractedContent(text=text, source=filename)
        except Exception as e:
            logger.error(f"解析字节数据失败 {filename}: {e}")


class CsvExtractor(BaseExtractor):
    """CSV提取器"""
    
    SUPPORTED_EXTENSIONS = [".csv", ".tsv"]
    
    def __init__(self, encoding: str = "utf-8", delimiter: Optional[str] = None):
        self.encoding = encoding
        self.delimiter = delimiter
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            with open(path, 'r', encoding=self.encoding, errors='ignore', newline='') as f:
                delimiter = self.delimiter or (',' if path.suffix == '.csv' else '\t')
                reader = csv.reader(f, delimiter=delimiter)
                for row_num, row in enumerate(reader):
                    text = ' '.join(str(cell) for cell in row)
                    yield ExtractedContent(
                        text=text,
                        source=path.name,
                        row=row_num,
                    )
        except Exception as e:
            logger.error(f"读取CSV失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            text = data.decode(self.encoding, errors='ignore')
            delimiter = self.delimiter or (',' if filename.endswith('.csv') else '\t')
            reader = csv.reader(StringIO(text), delimiter=delimiter)
            for row_num, row in enumerate(reader):
                row_text = ' '.join(str(cell) for cell in row)
                yield ExtractedContent(
                    text=row_text,
                    source=filename,
                    row=row_num,
                )
        except Exception as e:
            logger.error(f"解析CSV字节数据失败 {filename}: {e}")


class JsonExtractor(BaseExtractor):
    """JSON提取器"""
    
    SUPPORTED_EXTENSIONS = [".json", ".jsonl"]
    
    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
    
    def _extract_text_from_value(self, value, path: str = "") -> Generator[str, None, None]:
        """递归提取JSON中的文本值"""
        if isinstance(value, str):
            yield value
        elif isinstance(value, dict):
            for k, v in value.items():
                yield from self._extract_text_from_value(v, f"{path}.{k}")
        elif isinstance(value, list):
            for i, item in enumerate(value):
                yield from self._extract_text_from_value(item, f"{path}[{i}]")
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            with open(path, 'r', encoding=self.encoding, errors='ignore') as f:
                if path.suffix == '.jsonl':
                    # JSON Lines格式
                    for line_num, line in enumerate(f):
                        if line.strip():
                            try:
                                data = json.loads(line)
                                texts = list(self._extract_text_from_value(data))
                                yield ExtractedContent(
                                    text=' '.join(texts),
                                    source=path.name,
                                    row=line_num,
                                )
                            except json.JSONDecodeError:
                                continue
                else:
                    # 标准JSON
                    data = json.load(f)
                    texts = list(self._extract_text_from_value(data))
                    yield ExtractedContent(
                        text=' '.join(texts),
                        source=path.name,
                    )
        except Exception as e:
            logger.error(f"读取JSON失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            text = data.decode(self.encoding, errors='ignore')
            json_data = json.loads(text)
            texts = list(self._extract_text_from_value(json_data))
            yield ExtractedContent(
                text=' '.join(texts),
                source=filename,
            )
        except Exception as e:
            logger.error(f"解析JSON字节数据失败 {filename}: {e}")


class HtmlExtractor(BaseExtractor):
    """HTML提取器"""
    
    SUPPORTED_EXTENSIONS = [".html", ".htm", ".xhtml"]
    
    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding
    
    def _strip_tags(self, html: str) -> str:
        """简单的HTML标签移除"""
        import re
        # 移除script和style标签及其内容
        html = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        html = re.sub(r'<style[^>]*>.*?</style>', '', html, flags=re.DOTALL | re.IGNORECASE)
        # 移除所有HTML标签
        html = re.sub(r'<[^>]+>', ' ', html)
        # 处理HTML实体
        html = html.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
        # 压缩空白
        html = re.sub(r'\s+', ' ', html)
        return html.strip()
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            with open(path, 'r', encoding=self.encoding, errors='ignore') as f:
                html = f.read()
                text = self._strip_tags(html)
                yield ExtractedContent(text=text, source=path.name)
        except Exception as e:
            logger.error(f"读取HTML失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            html = data.decode(self.encoding, errors='ignore')
            text = self._strip_tags(html)
            yield ExtractedContent(text=text, source=filename)
        except Exception as e:
            logger.error(f"解析HTML字节数据失败 {filename}: {e}")


class PdfExtractor(BaseExtractor):
    """PDF提取器"""
    
    SUPPORTED_EXTENSIONS = [".pdf"]
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            import pypdf
            
            with open(path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        yield ExtractedContent(
                            text=text,
                            source=path.name,
                            page=page_num + 1,
                        )
        except ImportError:
            logger.warning("pypdf未安装，无法处理PDF文件。请运行: pip install pypdf")
        except Exception as e:
            logger.error(f"读取PDF失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            import pypdf
            
            reader = pypdf.PdfReader(BytesIO(data))
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    yield ExtractedContent(
                        text=text,
                        source=filename,
                        page=page_num + 1,
                    )
        except ImportError:
            logger.warning("pypdf未安装")
        except Exception as e:
            logger.error(f"解析PDF字节数据失败 {filename}: {e}")


class DocxExtractor(BaseExtractor):
    """Word文档提取器"""
    
    SUPPORTED_EXTENSIONS = [".docx"]
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            from docx import Document
            
            doc = Document(path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            yield ExtractedContent(
                text='\n'.join(paragraphs),
                source=path.name,
            )
        except ImportError:
            logger.warning("python-docx未安装，无法处理Word文件。请运行: pip install python-docx")
        except Exception as e:
            logger.error(f"读取Word文档失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            from docx import Document
            
            doc = Document(BytesIO(data))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            yield ExtractedContent(
                text='\n'.join(paragraphs),
                source=filename,
            )
        except ImportError:
            logger.warning("python-docx未安装")
        except Exception as e:
            logger.error(f"解析Word字节数据失败 {filename}: {e}")


class XlsxExtractor(BaseExtractor):
    """Excel提取器"""
    
    SUPPORTED_EXTENSIONS = [".xlsx", ".xls"]
    
    def extract(self, file_path: Union[str, Path]) -> Generator[ExtractedContent, None, None]:
        path = Path(file_path)
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    row_text = ' '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        yield ExtractedContent(
                            text=row_text,
                            source=f"{path.name}:{sheet_name}",
                            row=row_num,
                        )
            wb.close()
        except ImportError:
            logger.warning("openpyxl未安装，无法处理Excel文件。请运行: pip install openpyxl")
        except Exception as e:
            logger.error(f"读取Excel失败 {path}: {e}")
    
    def extract_from_bytes(self, data: bytes, filename: str) -> Generator[ExtractedContent, None, None]:
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    row_text = ' '.join(str(cell) if cell is not None else '' for cell in row)
                    if row_text.strip():
                        yield ExtractedContent(
                            text=row_text,
                            source=f"{filename}:{sheet_name}",
                            row=row_num,
                        )
            wb.close()
        except ImportError:
            logger.warning("openpyxl未安装")
        except Exception as e:
            logger.error(f"解析Excel字节数据失败 {filename}: {e}")


# 提取器注册表
EXTRACTORS: list[type[BaseExtractor]] = [
    TextExtractor,
    CsvExtractor,
    JsonExtractor,
    HtmlExtractor,
    PdfExtractor,
    DocxExtractor,
    XlsxExtractor,
]


def get_extractor(file_path: Union[str, Path]) -> Optional[BaseExtractor]:
    """
    根据文件类型获取对应的提取器
    
    :param file_path: 文件路径
    :return: 提取器实例，不支持的类型返回None
    """
    for extractor_class in EXTRACTORS:
        if extractor_class.supports(file_path):
            return extractor_class()
    return None
